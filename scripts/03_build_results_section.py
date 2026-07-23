from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
TABLES = ROOT / "outputs" / "tables"
FIGURES = ROOT / "outputs" / "figures"
OUTPUT = ROOT / "outputs" / "Results_section_English.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def format_table(table, widths):
    table.autofit = False
    table.style = "Table Grid"
    for r, row in enumerate(table.rows):
        for c, cell in enumerate(row.cells):
            cell.width = Inches(widths[c])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_mar = tc_pr.first_child_found_in("w:tcMar")
            if tc_mar is None:
                tc_mar = OxmlElement("w:tcMar")
                tc_pr.append(tc_mar)
            for side in ("top", "start", "bottom", "end"):
                node = tc_mar.find(qn(f"w:{side}"))
                if node is None:
                    node = OxmlElement(f"w:{side}")
                    tc_mar.append(node)
                node.set(qn("w:w"), "90")
                node.set(qn("w:type"), "dxa")
            for paragraph in cell.paragraphs:
                paragraph.alignment = (
                    WD_ALIGN_PARAGRAPH.LEFT if c == 0 else WD_ALIGN_PARAGRAPH.CENTER
                )
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.0
                for run in paragraph.runs:
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(8.5)
                    if r == 0:
                        run.bold = True
                        run.font.color.rgb = RGBColor(255, 255, 255)
            if r == 0:
                set_cell_shading(cell, "365F91")
            elif r % 2 == 0:
                set_cell_shading(cell, "EAF0F8")
    set_repeat_header(table.rows[0])


def caption(doc, label, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(7)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(f"{label}. {text}")
    r.bold = True
    r.font.name = "Times New Roman"
    r.font.size = Pt(9.5)


def add_figure(doc, letter, filename, legend, width=6.35):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_with_next = True
    p.add_run().add_picture(str(FIGURES / filename), width=Inches(width))
    caption(doc, f"Figure {letter}", legend)


def add_heading(doc, text, level=2):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    p.add_run(text)


def add_body(doc, text):
    p = doc.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Inches(0.3)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(6)
    return p


def add_dataframe_table(doc, df, headers, widths, formats=None):
    table = doc.add_table(rows=1, cols=len(headers))
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header
    for _, row in df.iterrows():
        cells = table.add_row().cells
        for i, key in enumerate(headers):
            value = row[key]
            formatter = (formats or {}).get(key)
            cells[i].text = formatter(value) if formatter else str(value)
    format_table(table, widths)
    return table


def main():
    totals = pd.read_csv(TABLES / "table_01_total_items_summary.csv")
    composition = pd.read_csv(TABLES / "table_02_category_composition.csv")
    cci = pd.read_csv(TABLES / "table_03_density_and_cci.csv")
    scores = pd.read_csv(TABLES / "table_04_pca_scores_and_clusters.csv")
    loadings = pd.read_csv(TABLES / "table_05_pca_loadings.csv")
    silhouette = pd.read_csv(TABLES / "table_06_cluster_selection.csv")
    one_factor = pd.read_csv(TABLES / "table_07_one_factor_permanova.csv")
    dispersion = pd.read_csv(TABLES / "table_08_permdisp.csv")
    factorial = pd.read_csv(TABLES / "table_09_factorial_permanova.csv")

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(11)
    normal.paragraph_format.line_spacing = 1.5

    for style_name, size in (("Heading 1", 14), ("Heading 2", 12)):
        style = doc.styles[style_name]
        style.font.name = "Times New Roman"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(0, 0, 0)
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(6)

    title = doc.add_paragraph(style="Heading 1")
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title.add_run("RESULTS")

    add_heading(doc, "Abundance and seasonal distribution")
    add_body(
        doc,
        "A total of 2,921 marine-litter items were recorded across the 24 "
        "transects. Praia Brava accounted for 2,677 items (91.6%), whereas "
        "Ilha do Pontal accounted for 244 items (8.4%). The largest seasonal "
        "total occurred in winter, when 2,183 items were recorded across both "
        "beaches (74.7% of the complete dataset). At Praia Brava, winter had "
        "the highest mean abundance (687.7 ± 564.3 items per transect), while "
        "autumn had the lowest (29.7 ± 28.2). At Ilha do Pontal, winter also "
        "had the highest mean (40.0 ± 24.3), whereas summer had the lowest "
        "(2.7 ± 2.5) (Table A; Figure A).",
    )

    t_a = totals.copy()
    t_a["Beach"] = t_a["beach"]
    t_a["Season"] = t_a["season"]
    t_a["n"] = t_a["n_transects"]
    t_a["Total"] = t_a["total_items"]
    t_a["Mean ± SD"] = t_a.apply(
        lambda r: f"{r['mean_items']:.1f} ± {r['sd_items']:.1f}", axis=1
    )
    t_a["Median (range)"] = t_a.apply(
        lambda r: f"{r['median_items']:.1f} ({int(r['minimum_items'])}–{int(r['maximum_items'])})",
        axis=1,
    )
    caption(
        doc,
        "Table A",
        "Seasonal abundance of marine-litter items at Praia Brava and Ilha do "
        "Pontal. Values summarize three transects per beach and season.",
    )
    add_dataframe_table(
        doc,
        t_a,
        ["Beach", "Season", "n", "Total", "Mean ± SD", "Median (range)"],
        [1.30, 0.85, 0.40, 0.65, 1.05, 1.35],
    )
    add_figure(
        doc,
        "A",
        "figure_01_total_items_by_season.png",
        "Marine-litter abundance by beach and season. Points and error bars "
        "represent transect-level summaries generated by the reproducible workflow.",
        width=5.5,
    )

    add_heading(doc, "Material composition")
    add_body(
        doc,
        "Expanded polystyrene was the most abundant category (1,432 items; "
        "49.0%), followed by plastic (1,059 items; 36.3%) and fragments "
        "(145 items; 5.0%). Together, expanded polystyrene and plastic "
        "represented 85.3% of all recorded items. Praia Brava contained "
        "1,395 expanded-polystyrene items and 899 plastic items, whereas Ilha "
        "do Pontal contained 37 and 160, respectively. The remaining categories "
        "each contributed less than 3% of the overall total (Table B; Figures B and C).",
    )
    t_b = composition.head(8).copy()
    remaining = composition.iloc[8:]
    other_row = {
        "category": "other_categories",
        "Praia Brava": int(remaining["Praia Brava"].sum()),
        "Ilha do Pontal": int(remaining["Ilha do Pontal"].sum()),
        "overall_total": int(remaining["overall_total"].sum()),
        "overall_percent": float(remaining["overall_percent"].sum()),
    }
    t_b = pd.concat([t_b, pd.DataFrame([other_row])], ignore_index=True)
    t_b["Category"] = t_b["category"].str.replace("_", " ").str.title()
    t_b["Praia Brava"] = t_b["Praia Brava"].astype(int)
    t_b["Ilha do Pontal"] = t_b["Ilha do Pontal"].astype(int)
    t_b["Overall"] = t_b["overall_total"].astype(int)
    t_b["Composition (%)"] = t_b["overall_percent"].map(lambda x: f"{x:.1f}")
    caption(
        doc,
        "Table B",
        "Composition of marine litter by category and beach, showing the eight "
        "most abundant categories; less abundant categories are pooled as Other.",
    )
    add_dataframe_table(
        doc,
        t_b,
        ["Category", "Praia Brava", "Ilha do Pontal", "Overall", "Composition (%)"],
        [1.75, 0.95, 1.05, 0.75, 1.10],
    )
    add_figure(
        doc,
        "B",
        "figure_02_composition_by_beach_and_season.png",
        "Relative composition of marine-litter categories by beach and season.",
    )
    add_figure(
        doc,
        "C",
        "figure_04_category_heatmap.png",
        "Heatmap of mean marine-litter abundance by category, beach, and season. "
        "Values are displayed on a log(1 + mean item count) scale.",
    )

    add_heading(doc, "Clean Coast Index")
    add_body(
        doc,
        "The Clean Coast Index (CCI) showed marked variation among beach-season "
        "combinations (Table C; Figure D). Praia Brava ranged from moderately "
        "dirty in autumn (CCI = 8.40) to extremely dirty in winter "
        "(CCI = 81.87); spring and summer were classified as dirty. Ilha do "
        "Pontal ranged from very clean in summer (CCI = 0.67) to moderately "
        "dirty in autumn (CCI = 7.07) and winter (CCI = 9.07), while spring "
        "was classified as clean (CCI = 4.53).",
    )
    t_c = cci.copy()
    t_c["Beach"] = t_c["beach"]
    t_c["Season"] = t_c["season"]
    t_c["Items m⁻²"] = t_c["total_density_items_m2"].map(lambda x: f"{x:.2f}")
    t_c["Plastic items m⁻²"] = t_c["plastic_density_items_m2"].map(
        lambda x: f"{x:.2f}"
    )
    t_c["CCI"] = t_c["cci"].map(lambda x: f"{x:.2f}")
    t_c["Classification"] = t_c["cci_classification"]
    caption(
        doc,
        "Table C",
        "Marine-litter density, plastic density, and Clean Coast Index (CCI) "
        "classification by beach and season.",
    )
    add_dataframe_table(
        doc,
        t_c,
        ["Beach", "Season", "Items m⁻²", "Plastic items m⁻²", "CCI", "Classification"],
        [1.15, 0.75, 0.80, 1.05, 0.55, 1.15],
    )
    add_figure(
        doc,
        "D",
        "figure_05_clean_coast_index.png",
        "Clean Coast Index (CCI) by beach and season. Dashed horizontal lines "
        "indicate classification thresholds.",
    )

    add_heading(doc, "Multivariate structure and clustering")
    top_loadings = loadings.sort_values("contribution_percent", ascending=False).head(6)
    add_body(
        doc,
        "The first two axes of the Hellinger PCA explained 63.95% of the "
        "variation in litter composition (PC1 = 41.19%; PC2 = 22.76%). "
        "Expanded polystyrene made the largest contribution to the ordination "
        "(41.28%), followed by plastic (28.32%) and processed wood (17.61%). "
        "Expanded polystyrene loaded positively on PC1, plastic loaded strongly "
        "and positively on PC2, and processed wood loaded negatively on PC2 "
        "(Table D; Figure E).",
    )
    t_d = top_loadings.copy()
    t_d["Category"] = t_d["category"].str.replace("_", " ").str.title()
    t_d["PC1 loading"] = t_d["PC1_loading"].map(lambda x: f"{x:.3f}")
    t_d["PC2 loading"] = t_d["PC2_loading"].map(lambda x: f"{x:.3f}")
    t_d["Contribution (%)"] = t_d["contribution_percent"].map(lambda x: f"{x:.2f}")
    caption(
        doc,
        "Table D",
        "Largest category contributions to the first two axes of the "
        "Hellinger-transformed principal component analysis.",
    )
    add_dataframe_table(
        doc,
        t_d,
        ["Category", "PC1 loading", "PC2 loading", "Contribution (%)"],
        [2.05, 1.05, 1.05, 1.20],
    )
    add_figure(
        doc,
        "E",
        "figure_03_hellinger_pca.png",
        "Hellinger-transformed principal component analysis of marine-litter "
        "composition. Ellipses represent 68% group regions by beach; vectors "
        "show category loadings and are colored by contribution.",
    )

    best = silhouette.loc[silhouette["silhouette_score"].idxmax()]
    counts = scores.groupby("cluster").size().to_dict()
    add_body(
        doc,
        "K-means clustering selected three clusters because k = 3 produced the "
        f"highest silhouette score ({best['silhouette_score']:.3f}; Table E). "
        f"Cluster 1 contained {counts.get(1, 0)} samples, cluster 2 contained "
        f"{counts.get(2, 0)} sample, and cluster 3 contained {counts.get(3, 0)} "
        "samples (Figure F). Because cluster 2 contained a single sample, a "
        "within-cluster confidence ellipse could not be estimated for that group.",
    )
    t_e = silhouette.copy()
    t_e["Number of clusters (k)"] = t_e["k"].astype(int)
    t_e["Silhouette score"] = t_e["silhouette_score"].map(lambda x: f"{x:.3f}")
    caption(
        doc,
        "Table E",
        "Silhouette scores used to select the number of K-means clusters.",
    )
    add_dataframe_table(
        doc,
        t_e,
        ["Number of clusters (k)", "Silhouette score"],
        [2.7, 2.7],
    )
    add_figure(
        doc,
        "F",
        "figure_06_kmeans_clusters.png",
        "K-means clusters projected onto the first two Hellinger PCA axes. "
        "Colors and symbols identify clusters, enlarged symbols indicate "
        "centroids, and ellipses show 68% within-cluster regions when estimable.",
    )

    add_heading(doc, "Differences in multivariate composition")
    add_body(
        doc,
        "The one-factor PERMANOVA detected a difference in composition between "
        "beaches (pseudo-F = 3.548, p = 0.0085), whereas the season-only test "
        "was not significant (pseudo-F = 1.195, p = 0.2744). PERMDISP did not "
        "detect differences in multivariate dispersion for beach (F = 0.309, "
        "p = 0.5623) or season (F = 0.891, p = 0.3441). In the factorial "
        "PERMANOVA, beach was significant (pseudo-F = 4.929, R² = 0.139, "
        "p = 0.0020), season was marginal but not significant at α = 0.05 "
        "(pseudo-F = 1.798, R² = 0.152, p = 0.0523), and the beach × season "
        "interaction was significant (pseudo-F = 3.058, R² = 0.258, "
        "p = 0.0014) (Table F).",
    )
    rows = []
    for _, r in one_factor.iterrows():
        rows.append(
            {
                "Analysis": "PERMANOVA",
                "Term": r["factor"].title(),
                "Statistic": f"pseudo-F = {r['pseudo_f']:.3f}",
                "R²": "—",
                "p": f"{r['p_value']:.4f}",
            }
        )
    for _, r in dispersion.iterrows():
        rows.append(
            {
                "Analysis": "PERMDISP",
                "Term": r["factor"].title(),
                "Statistic": f"F = {r['f_value']:.3f}",
                "R²": "—",
                "p": f"{r['p_value']:.4f}",
            }
        )
    for _, r in factorial.iterrows():
        rows.append(
            {
                "Analysis": "Factorial PERMANOVA",
                "Term": r["term"].replace(":", " × ").title(),
                "Statistic": f"pseudo-F = {r['pseudo_f']:.3f}",
                "R²": f"{r['r_squared']:.3f}",
                "p": f"{r['p_value']:.4f}",
            }
        )
    t_f = pd.DataFrame(rows)
    caption(
        doc,
        "Table F",
        "Multivariate tests of marine-litter composition. All permutation tests "
        "used 9,999 permutations.",
    )
    add_dataframe_table(
        doc,
        t_f,
        ["Analysis", "Term", "Statistic", "R²", "p"],
        [1.45, 1.25, 1.40, 0.60, 0.70],
    )

    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
